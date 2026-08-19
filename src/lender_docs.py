"""The lender document set (build 1.0.42, 2026-08).

Seven Word documents that ship alongside the workbook. Together they are the
8-piece set the cover letter describes:

    1. Cover Letter
    2. The Private Lender Package        <- lender_package.py builds this
    3. Promissory Note
    4. Personal Guarantee
    5. Closing Instructions Letter
    6. Insurance Request Letter
    7. Investor Information Sheet
    8. Satisfaction and Release Request

These are transcriptions of the templates the team edited by hand, with the
highlighted fields filled from the deal spec. Rules taken from their notes:

  - Highlighted text in the source templates marks a field to replace, so
    anything still unresolved renders as a visible [PLACEHOLDER] rather than
    quietly disappearing.
  - Where the template offers two options separated by OR, exactly one gets
    written. `deal_type` picks wholetail or flip; the presence of a repair
    budget picks the insurance policy type.
  - The NOTES FOR CLAUDE block never appears in any output. There is an
    assertion at the end of the build that proves it.
  - Every member of the company guarantees the note, so the guarantee and the
    closing letter render one block per member.

Usage:
    python src/lender_docs.py --spec deals/158_old_state_lender.json
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

NAVY = RGBColor(0x0A, 0x11, 0x30)
GREY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xC0, 0x00, 0x00)

DISCLAIMER = (
    "This document is a template and a starting point for your attorney. It is not "
    "legal advice and it is not a substitute for legal advice. Laws vary from state "
    "to state and change over time. Both the lender and the borrower should have "
    "counsel review the executed loan documents before signing."
)
BANNED = "NOTES FOR CLAUDE"


def _money(v) -> str:
    if not isinstance(v, (int, float)):
        return str(v or "")
    return f"-${abs(v):,.0f}" if v < 0 else f"${v:,.0f}"


def _words(n: int) -> str:
    ones = ["zero", "one", "two", "three", "four", "five", "six", "seven", "eight",
            "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen",
            "sixteen", "seventeen", "eighteen", "nineteen"]
    tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy",
            "eighty", "ninety"]

    def u1000(x):
        if x < 20:
            return ones[x]
        if x < 100:
            return tens[x // 10] + ("-" + ones[x % 10] if x % 10 else "")
        return ones[x // 100] + " hundred" + (" " + u1000(x % 100) if x % 100 else "")

    n = int(round(n))
    if n == 0:
        return "zero"
    parts = []
    for div, name in ((1_000_000, "million"), (1_000, "thousand")):
        if n >= div:
            parts.append(u1000(n // div) + " " + name)
            n %= div
    if n:
        parts.append(u1000(n))
    return " ".join(parts)


# ── docx helpers ───────────────────────────────────────────────────────

def _doc():
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(8)
    return d


def _h1(d, text, sub=""):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    if sub:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sub.upper()); r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = NAVY
    return d


def _h2(d, text):
    p = d.add_paragraph()
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    return p


def _p(d, text, bold=False, color=None, size=10.5, align=None):
    p = d.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def _bullet(d, text):
    p = d.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(10.5)
    return p


def _field(d, label, value):
    p = d.add_paragraph()
    r = p.add_run(f"{label}: "); r.bold = True; r.font.size = Pt(10.5)
    r = p.add_run(str(value)); r.font.size = Pt(10.5)
    return p


def _sig(d, role, name):
    d.add_paragraph()
    _p(d, role.upper(), bold=True, size=9.5, color=NAVY)
    _p(d, "_" * 54)
    _p(d, name, size=9.5)
    _p(d, "Date: ______________________", size=9.5)


def _disclaim(d):
    d.add_paragraph()
    _p(d, DISCLAIMER, color=GREY, size=8.5)


def _ctx(spec):
    from lender_package import compute
    m = compute(spec)
    b, p, ln = spec["borrower"], spec["property"], spec["loan"]
    members = b.get("members") or [b.get("signer") or "[SIGNER NAME]"]
    return {
        "m": m, "b": b, "p": p, "ln": ln, "members": members,
        "entity": b["entity"],
        "formation": b.get("formation", "a Tennessee limited liability company"),
        "signer": b.get("signer") or "[SIGNER NAME]",
        "lender": spec.get("lender_name") or "[LENDER LEGAL NAME]",
        "date": spec.get("date", "[DATE]"),
        "agent": spec.get("closing_agent") or "[CLOSING ATTORNEY / TITLE COMPANY]",
        "county": p.get("county", "Knox"),
        "is_flip": spec.get("deal_type", "flip").lower() == "flip",
    }


# ── 1. Cover Letter ────────────────────────────────────────────────────

def doc_cover(spec, out: Path):
    c = _ctx(spec); m, p, b = c["m"], c["p"], c["b"]
    d = _doc()
    _h1(d, "Private lending opportunity", p["full_address"])
    d.add_paragraph()
    _p(d, c["date"], color=GREY, size=9.5)
    _p(d, f"{spec.get('lender_salutation', '[LENDER FIRST NAME]')},")
    _p(d, "We have a deal that we believe would be a great fit, and we want you to "
          "fund on it. Here are the details:")

    if c["is_flip"]:
        _p(d, "Deal Type. We would be doing a flip. This is the best way to guarantee "
              "you make the most on your return. The goal of this strategy is to find a "
              "property and fully renovate it before listing on the MLS. The time needed "
              "to perform the renovation is what secures a bigger pay day for you.")
        _p(d, f"Purchase Price. We're buying a house at "
              f"{_money(spec['contract_price'])}, "
              f"renovating it, and plan to relist it at {_money(m['arv'])} depending on "
              f"comps at that time. The only thing that has to happen for you to get paid "
              f"back is that the house sells on the open market for more than "
              f"{_money(m['payoff'])}. We believe this is very achievable based on the "
              f"comps included in the breakdown provided, before appreciation increasing "
              f"home prices over the months we hold the property.")
    else:
        _p(d, "Deal Type. We would be doing a wholetail. What that means is we are "
              "looking to purchase the property, and then re-list it within 30 days. "
              "This is the simplest type of deal for an investor to run, with a turn "
              "around time much faster than a normal flip because there is no renovation "
              "needed.")
        _p(d, f"Purchase Price. We're buying a house at "
              f"{_money(spec['contract_price'])} and relisting it on the MLS in exactly "
              f"the condition we bought it. The only thing that has to happen for you to "
              f"get paid back is that the house sells on the open market for more than "
              f"{_money(m['payoff'])}. We believe this is very achievable based on the "
              f"comps included in the breakdown provided.")

    _h2(d, "The ask, in one line")
    _p(d, f"{_money(m['loan'])} at a {m['rate']*100:g}% annualized rate for "
          f"{m['term']} months, with {m['min_int_pct']*100:.1f}% guaranteed, secured in "
          f"first position by a deed of trust on {p['short_address']}. You get "
          f"{_money(m['payoff'])} back in one payment out of the closing when it sells, "
          f"and you get paid before we see a dollar.", bold=True)

    _h2(d, "Why your money is safe here")
    for line in spec["plan"].get("safety_bullets", []):
        _bullet(d, line)
    _bullet(d, f"{m['term']} month term on a plan that expects to be done in "
               f"{spec['plan'].get('work_window','less')}, and you earn "
               f"{m['min_int_months']} months of interest minimum even if it sells in "
               f"thirty days.")
    _bullet(d, f"We include a personal guarantee. Any property that we default on, we "
               f"immediately look to liquidate the property for you to recover your "
               f"funds, and the members of {c['entity']} guarantee to make up the "
               f"remaining amount not captured from the sale including the full loan "
               f"amount plus the interest on the loan.")

    _h2(d, "What is included")
    for name, desc in [
        ("The Private Lender Package",
         "This is a deep dive into the stats of the deal itself. It includes an "
         "overview of the deal itself, an overview of your contribution to the deal, a "
         "term sheet, an in-depth dive into the numbers on repairs (if available) and "
         "re-sell value, the comps that support the deal logic, back ups and risk "
         "mitigation in the worst case scenario, and next steps for you to fill out."),
        ("The Promissory Note",
         f"This gives a clear breakdown of the exact loan terms, which will be signed "
         f"by us at {c['entity']} to give you the right to default if any of the "
         f"instances stated in this document occur."),
        ("The Personal Guarantee",
         f"This document is to ensure that no matter what happens with the deal, you "
         f"will get paid back the full amount owed plus interest. The process is to "
         f"move to immediately liquidate the property for the max amount available on "
         f"the market, and then pay back the rest through the personal guarantee split "
         f"across all members of {c['entity']}."),
        ("The Closing Instructions Letter",
         "This is the document prepared for title to understand our full situation, so "
         "they can be expecting your funds and make sure your investment is not "
         "mishandled in any way."),
        ("The Insurance Request Letter",
         "A letter to the temporary insurance company putting the policy under your "
         "name, so if there is any issue in regards to the property and the insurance "
         "company reimburses us, they are reimbursing you and not us."),
        ("The Investor Information Sheet",
         "This sheet is filled out after close for your benefit. You can see the exact "
         "details of the loan, term, insurance, and your rights as an investor."),
        ("The Satisfaction and Release Request",
         "This document is executed at the time of the resell, to confirm the loan was "
         "repaid and you have your money back."),
    ]:
        p_ = d.add_paragraph(style="List Bullet")
        r = p_.add_run(f"{name}: "); r.bold = True; r.font.size = Pt(10.5)
        r = p_.add_run(desc); r.font.size = Pt(10.5)

    _h2(d, "What to do next")
    _p(d, "Read it, poke holes in it, and call me with whatever does not add up. If the "
          "answer is no, tell me no and it will not change anything between us. I would "
          "rather you passed on this one and funded the next one than felt pushed into "
          "either.")
    _p(d, "If you are wanting to go forward with the deal, please go to The Private "
          "Lender Package and fill in the highlighted cells under the Next Steps tab.")
    d.add_paragraph()
    _p(d, c["signer"], bold=True)
    _p(d, f"Managing Member, {c['entity']}", size=9.5, color=GREY)
    _p(d, f"{b.get('phone') or '[PHONE]'} | {b.get('email') or '[EMAIL]'}",
       size=9.5, color=GREY)
    d.add_paragraph()
    _p(d, "One housekeeping note that matters more than it sounds like it does. When it "
          "comes time to fund, call the closing attorney's office at a number you looked "
          "up yourself and confirm the wire instructions out loud. Never work off wire "
          "instructions that only showed up in an email, including one from me.",
       color=RED, size=9.5)
    d.save(out)


# ── 3. Promissory Note ─────────────────────────────────────────────────

def doc_note(spec, out: Path):
    c = _ctx(spec); m, p, ln = c["m"], c["p"], c["ln"]
    d = _doc()
    _h1(d, "Promissory Note")
    _p(d, f"Principal Amount: {_money(m['loan'])}          Date: {c['date']}          "
          f"Place: {c['county']} County, Tennessee",
       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph()
    _p(d, f"FOR VALUE RECEIVED, the undersigned, {c['entity']}, {c['formation']} "
          f"(\"Borrower\"), promises to pay to the order of {c['lender']} (\"Lender\"), "
          f"at such address as Lender may designate in writing, the principal sum of "
          f"{_words(m['loan']).upper()} DOLLARS ({_money(m['loan'])}), together with "
          f"interest on the unpaid principal balance at the rate of "
          f"{m['rate']*100:g} percent ({m['rate']*100:g}%) per annum, simple interest, "
          f"from the date funds are advanced until paid in full.")

    _h2(d, "1. Security")
    _p(d, f"This Note is secured by a Deed of Trust of even date granting Lender a first "
          f"lien on the real property commonly known as {p['full_address']}, "
          f"{c['county']} County, Tennessee (the \"Property\"), to be recorded in the "
          f"Register of Deeds of that county.")

    _h2(d, "2. Payment")
    _p(d, f"No monthly payments of principal or interest are required. The entire unpaid "
          f"principal balance together with all accrued and unpaid interest shall be due "
          f"and payable in full on the earlier of (a) the closing of a sale or transfer "
          f"of the Property, or (b) {m['term']} months from the date funds are advanced "
          f"(the \"Maturity Date\").")
    _p(d, f"Minimum interest. Notwithstanding any earlier payoff, Borrower shall pay "
          f"Lender not less than {m['min_int_months']} months of interest on the "
          f"principal advanced, which is {m['min_int_pct']*100:.1f}% of the principal "
          f"({_money(m['min_int_amt'])}).")
    _p(d, "Prepayment. Borrower may prepay all or any part of the principal at any time "
          "without penalty, subject only to the minimum interest above. Payments shall "
          "be applied first to accrued interest and the balance to principal.")

    _h2(d, "3. Advance")
    if m["draws"]:
        _p(d, f"{_money(m['day_one'])} of the principal sum shall be advanced at closing "
              f"through the closing agent's escrow account, and the balance of "
              f"{_money(m['rehab_tranche'])} shall be advanced in {m['draws']} draws "
              f"upon completion of the corresponding phase of work, on Borrower's "
              f"written request supported by photographs, paid invoices and lien "
              f"waivers. Interest accrues only on amounts actually advanced.")
    else:
        _p(d, "The entire principal sum shall be advanced in a single disbursement at "
              "closing through the closing agent's escrow account.")

    _h2(d, "4. Insurance")
    _p(d, "Borrower shall keep the Property insured at all times under a policy "
          "appropriate to its occupancy and condition, naming Lender as mortgagee and "
          "loss payee, and shall deliver evidence of that coverage to Lender before any "
          "funds are advanced and upon each renewal. Failure to maintain such insurance "
          "is an event of default.")

    _h2(d, "5. Default")
    _p(d, "Each of the following is an event of default: failure to pay any amount when "
          "due; failure to maintain insurance as required; sale, transfer or further "
          "encumbrance of the Property without Lender's prior written consent; or any "
          "material misrepresentation made to Lender in connection with this loan.")
    _p(d, "Upon an event of default that remains uncured for ten (10) days after written "
          "notice from Lender, the entire unpaid principal balance and all accrued "
          "interest shall, at the option of Lender, become immediately due and payable. "
          "The Borrower shall immediately move to liquidate the property for as much as "
          "possible on the market, and the personal guarantee will take effect to cover "
          "any amount still owed to the Lender including the amount of interest still "
          "left owed.")

    _h2(d, "6. Collection Costs")
    _p(d, "Borrower agrees to pay all costs of collection, including reasonable "
          "attorneys' fees and court costs, incurred by Lender in enforcing this Note, "
          "to the extent permitted by law.")

    _h2(d, "7. Guaranty")
    _p(d, f"Payment of this Note is personally guaranteed by all members of "
          f"{c['entity']} under a separate Guaranty Agreement of even date.")

    _h2(d, "8. Waivers and Miscellaneous")
    _p(d, "Borrower and all endorsers, guarantors and sureties waive demand, "
          "presentment, protest, notice of dishonor and notice of protest. Lender's "
          "failure to exercise any right is not a waiver of that right in the future. No "
          "modification of this Note is binding unless in writing and signed by both "
          "parties. If any provision is held unenforceable, the remainder stays in full "
          "force. This Note shall be governed by and construed in accordance with the "
          "laws of the State of Tennessee.")

    d.add_paragraph()
    _p(d, "BORROWER:", bold=True)
    _p(d, f"{c['entity']}, {c['formation']}")
    _sig(d, "By", f"{c['signer']}, {c['b'].get('title','Managing Member')}")
    d.add_paragraph()
    _p(d, "Instructions: sign one original only. Lender keeps the original with the "
          "original signature; Borrower keeps a copy. At payoff, Lender writes PAID IN "
          "FULL across the face of the original, signs and dates it, and returns it to "
          "Borrower.", color=GREY, size=9)
    _disclaim(d)
    d.save(out)


# ── 4. Personal Guarantee ──────────────────────────────────────────────

def doc_guarantee(spec, out: Path):
    c = _ctx(spec); m, p = c["m"], c["p"]
    d = _doc()
    _h1(d, "Guaranty Agreement")
    _p(d, f"Date: {c['date']}", bold=True)
    d.add_paragraph()
    _p(d, f"In consideration of {c['lender']} (\"Lender\") making a loan in the "
          f"principal amount of {_money(m['loan'])} to {c['entity']} (\"Borrower\"), "
          f"which loan is evidenced by a Promissory Note of even date and secured by a "
          f"Deed of Trust on {p['full_address']}, the undersigned "
          f"{', '.join(c['members'])} (\"Guarantors\"), individually, absolutely and "
          f"unconditionally guarantee to Lender the full and prompt payment and "
          f"performance of all obligations of Borrower under the Note.")

    _h2(d, "1. Nature of the Guaranty")
    _p(d, "This is a guaranty of payment and not merely of collection. Lender is not "
          "required to first proceed against Borrower, against the Property, or against "
          "any other security before enforcing this Guaranty.")

    _h2(d, "2. Continuing Obligation")
    _p(d, "This Guaranty remains in effect until the Note is paid in full. Guarantors' "
          "obligations are not released or reduced by any extension, modification, "
          "renewal or indulgence granted to Borrower, by the release of any collateral, "
          "or by any failure or delay by Lender in enforcing any right.")

    _h2(d, "3. Waivers")
    _p(d, "Guarantors waive demand, presentment, protest, notice of dishonor, notice of "
          "acceptance of this Guaranty and notice of any default by Borrower.")

    _h2(d, "4. Costs")
    _p(d, "Guarantors agree to pay Lender's reasonable attorneys' fees and costs "
          "incurred in enforcing this Guaranty, to the extent permitted by law.")

    _h2(d, "5. Governing Law")
    _p(d, "This Guaranty shall be governed by and construed in accordance with the laws "
          "of the State of Tennessee.")

    for member in c["members"]:
        _sig(d, "Guarantor", f"{member}, individually")
        _p(d, f"Guarantor address: {c['b'].get('address') or '[BUSINESS ADDRESS]'}",
           size=9.5)
    _disclaim(d)
    d.save(out)


# ── 5. Closing Instructions Letter ─────────────────────────────────────

def doc_closing(spec, out: Path):
    c = _ctx(spec); m, p, ln = c["m"], c["p"], c["ln"]
    d = _doc()
    _h1(d, "Closing Instruction Letter")
    _field(d, "To", c["agent"])
    _field(d, "From", c["entity"])
    _field(d, "Date", c["date"])
    _field(d, "Re", f"Private money loan secured by {p['full_address']}")
    d.add_paragraph()
    _p(d, "Please prepare the following file for closing. This is a private money loan "
          "from an individual lender, not an institutional loan, and the lender is "
          "relying on you to protect their lien position.")

    _h2(d, "Parties")
    _field(d, "Lender", c["lender"])
    _field(d, "Borrower", f"{c['entity']}, {c['formation']}")
    for member in c["members"]:
        _field(d, "Guarantor", f"{member}, individually")
    _field(d, "Seller", spec.get("seller_name") or "[SELLER NAME]")
    _field(d, "Property", f"{p['full_address']}, parcel {p.get('parcel_id','')}")

    _h2(d, "Loan terms to document")
    _field(d, "Loan amount", _money(m["loan"]))
    _field(d, "Interest rate", f"{m['rate']*100:g}% per annum, simple")
    _field(d, "Points and fees", f"{m['points']:g} points ({_money(m['points_d'])})")
    _field(d, "Term", f"{m['term']} months from funding")
    _field(d, "Payment", "No monthly payments; principal and accrued interest due at "
                         "the earlier of sale or maturity")
    _field(d, "Minimum interest", f"{m['min_int_months']} months "
                                  f"({m['min_int_pct']*100:.1f}% of the loan, "
                                  f"{_money(m['min_int_amt'])}) regardless of earlier "
                                  f"payoff")
    _field(d, "Security", f"{ln.get('lien','First lien, deed of trust')} on the Property")
    _field(d, "Disbursement",
           "Single advance at closing" if not m["draws"] else
           f"{_money(m['day_one'])} at closing, balance of "
           f"{_money(m['rehab_tranche'])} released in {m['draws']} draws against "
           f"completed work")
    _field(d, "Default", "Full amount owed including interest at point of default, with "
                         "immediate liquidation of the property and the personal "
                         "guarantee covering any shortfall")
    _field(d, "Governing law", "Tennessee")

    _h2(d, "Documents to prepare")
    for i, item in enumerate([
        f"Deed of Trust on your standard Tennessee form, granting Lender a first lien "
        f"on the Property, in recordable form, to be recorded with the {c['county']} "
        f"County Register of Deeds on closing day.",
        "Lender's title insurance policy in favor of Lender in the amount of the loan, "
        "with no exceptions other than those approved in writing by Lender.",
        "Settlement statement showing the loan proceeds, the purchase, and all costs.",
        "Confirmation to Lender in writing when their wired funds are received, before "
        "any disbursement occurs.",
    ], 1):
        _p(d, f"{i}. {item}")
    _p(d, "The Promissory Note and Guaranty Agreement are attached and executed by "
          "Borrower and Guarantors at closing. Borrower is providing them; you do not "
          "need to draft them, but please review them for consistency with the Deed of "
          "Trust you prepare.")

    _h2(d, "Conditions that must be satisfied before disbursement")
    for i, item in enumerate([
        "Title commitment clear and Lender's policy issued in the loan amount.",
        "Hazard or builder's risk insurance in force on the Property naming Lender as "
        "mortgagee and loss payee, certificate in your file.",
        "Deed of Trust executed and in recordable form.",
        "Promissory Note and Guaranty executed.",
        "Lender's funds received into your escrow account and confirmed.",
    ], 1):
        _p(d, f"{i}. {item}")

    _h2(d, "Wire instructions")
    _p(d, "Please send wire instructions to the Lender DIRECTLY from your office, and "
          "expect a phone call from the Lender to verbally confirm them before they send "
          "funds. Do not accept changed wire instructions by email from anyone, "
          "including from us.", bold=True, color=RED)

    _h2(d, "Costs")
    _p(d, "The lender's loan amount covers all closing costs, document preparation, "
          "recording fees and the lender's title policy premium, to be paid back with "
          "interest and backed by the personal guarantee.")

    _sig(d, "Borrower", f"{c['signer']}, {c['b'].get('title','Managing Member')}, "
                        f"{c['entity']}")
    _disclaim(d)
    d.save(out)


# ── 6. Insurance Request Letter ────────────────────────────────────────

def doc_insurance(spec, out: Path):
    c = _ctx(spec); m, p = c["m"], c["p"]
    d = _doc()
    _h1(d, "Insurance Request Letter")
    _field(d, "To", spec.get("insurance_agent") or "[INSURANCE AGENT]")
    _field(d, "From", c["entity"])
    _field(d, "Date", c["date"])
    _field(d, "Re", f"Coverage required at closing, {p['full_address']}")
    d.add_paragraph()
    _p(d, "We are closing on the property below and need coverage bound and evidenced "
          "BEFORE our lender's funds disburse. Please issue a certificate to the closing "
          "attorney and a copy to the lender named below.")

    _h2(d, "Property")
    _field(d, "Address", p["full_address"])
    _field(d, "Parcel", p.get("parcel_id", ""))
    _field(d, "Year built", p.get("year_built", ""))
    _field(d, "Square footage", f"{p.get('sqft',0):,}")
    _field(d, "Occupancy", p.get("occupancy", "Vacant"))
    _field(d, "Construction", p.get("type", ""))
    _field(d, "Condition", p.get("condition", ""))

    _h2(d, "Coverage required")
    # The template offers two policy types separated by OR. A repair budget
    # means there is a renovation to insure, so the deal picks the branch.
    _field(d, "Policy type",
           "Renovation plan available, builder's risk" if m["rehab"] else
           "Vacant dwelling, no renovation planned")
    _field(d, "Dwelling limit", f"Not less than {_money(m['loan'])} (the loan amount)")
    _field(d, "Liability", "$1,000,000 general liability")
    _field(d, "Term", f"{m['term']} months with a short rate cancellation")
    _field(d, "Effective date", "On or before the closing date, no gap")

    _h2(d, "Mortgagee and loss payee clause")
    _p(d, "Please name the following exactly as written, as MORTGAGEE and LOSS PAYEE:",
       bold=True)
    _p(d, c["lender"], bold=True)
    _p(d, spec.get("lender_address") or "[LENDER MAILING ADDRESS]")
    d.add_paragraph()
    _p(d, "The lender must be notified directly by the carrier if the policy lapses, is "
          "cancelled or is not renewed. Please confirm that the notification is set up.")

    _h2(d, "Deliver to")
    _field(d, "Closing attorney", c["agent"])
    _field(d, "Lender copy", c["lender"])
    _field(d, "Borrower copy", c["b"].get("email") or "[OUR EMAIL]")
    _disclaim(d)
    d.save(out)


# ── 7. Investor Information Sheet ──────────────────────────────────────

def doc_investor(spec, out: Path):
    c = _ctx(spec); m, p = c["m"], c["p"]
    d = _doc()
    _h1(d, "Investor Information Sheet", c["entity"])
    _p(d, "Kept in the lender file for this loan. One sheet per lender per property.",
       color=GREY, size=9)

    _h2(d, "Lender")
    for f in ["Name exactly as it appears on the Note", "Entity, trust or individual",
              "Mailing address", "City, state, ZIP", "Phone", "Email",
              "Custodian (if lending from a retirement account)",
              "Custodian account number", "Lender's attorney (optional)"]:
        _p(d, f"{f}: {'_' * max(8, 68 - len(f))}")

    _h2(d, "Terms of investment")
    _field(d, "Property", p["full_address"])
    _field(d, "Amount invested", _money(m["loan"]))
    _field(d, "Date of investment", "______________________")
    _field(d, "Interest rate", f"{m['rate']*100:g}% per annum")
    _field(d, "Points", f"{m['points']:g} ({_money(m['points_d'])})")
    _field(d, "Term", f"{m['term']} months")
    _field(d, "Maturity date", "______________________")
    _field(d, "Payment schedule", "Principal and accrued interest at payoff")
    _field(d, "Minimum interest", f"{m['min_int_months']} months "
                                  f"({_money(m['min_int_amt'])})")
    _field(d, "Estimated payoff", _money(m["payoff"]))
    _field(d, "Secured by", "First lien, deed of trust, plus a personal guarantee from "
                            "every member of the company")
    _field(d, "Recorded", f"{c['county']} County Register of Deeds, Book ____ "
                          f"Page ____ Instrument ____")
    _field(d, "Title policy number", "______________________")
    _field(d, "Insurance carrier and policy number", "______________________")

    _h2(d, "Communication preferences")
    for f in ["Update frequency (default: every two weeks with photographs)",
              "Preferred contact method", "Wants site visits (yes / no)",
              "Interested in funding future deals (yes / no)",
              "Maximum they would deploy on one deal"]:
        _p(d, f"{f}: {'_' * max(8, 66 - len(f))}")

    _h2(d, "Notes")
    for _ in range(6):
        _p(d, "_" * 100)
    _disclaim(d)
    d.save(out)


# ── 8. Satisfaction and Release Request ────────────────────────────────

def doc_release(spec, out: Path):
    c = _ctx(spec); m, p = c["m"], c["p"]
    d = _doc()
    _h1(d, "Satisfaction and Release Request")
    _p(d, "Used at payoff. Send to the lender with the payoff figure once the resale "
          "closing is scheduled.", color=GREY, size=9)
    d.add_paragraph()
    _field(d, "To", c["lender"])
    _field(d, "From", c["entity"])
    _field(d, "Date", "______________________")
    _field(d, "Re", f"Payoff and release, {p['full_address']}")
    d.add_paragraph()
    _p(d, "The sale of the above property is scheduled to close on ______________. Your "
          "loan will be paid in full directly from the settlement, before any proceeds "
          "are released to us.")

    _h2(d, "Payoff calculation")
    _field(d, "Original principal", _money(m["loan"]))
    _field(d, "Date funds advanced", "______________________")
    _field(d, "Payoff date", "______________________")
    _field(d, "Days outstanding", "______________________")
    _field(d, "Interest accrued at " + f"{m['rate']*100:g}% per annum",
           "______________________")
    _field(d, f"Minimum interest floor ({m['min_int_months']} months)",
           _money(m["min_int_amt"]))
    _field(d, "TOTAL PAYOFF", "______________________")
    _field(d, "Wire to", "Lender's account, instructions to be confirmed by phone")

    _h2(d, "On receipt of payoff, please")
    for i, item in enumerate([
        "Write PAID IN FULL across the face of the original Promissory Note, sign and "
        "date it, and return the original to us.",
        f"Execute the Release or Satisfaction of Deed of Trust prepared by the closing "
        f"attorney so it can be recorded with the {c['county']} County Register of Deeds.",
        "Confirm in writing that the obligation is satisfied and that you have no "
        "further claim against the property.",
    ], 1):
        _p(d, f"{i}. {item}")

    d.add_paragraph()
    _p(d, "It has been a pleasure working with you. If you would like your funds placed "
          "on another property, tell us now and we will hold your position on the next "
          "one that clears our buying criteria.")
    _sig(d, "Borrower", f"{c['signer']}, {c['b'].get('title','Managing Member')}, "
                        f"{c['entity']}")
    d.add_paragraph()
    _p(d, "LENDER ACKNOWLEDGMENT OF PAYMENT IN FULL", bold=True, color=NAVY)
    _sig(d, "Lender", c["lender"])
    _disclaim(d)
    d.save(out)


# Numbering matches the cover letter. 2 is the workbook, built elsewhere.
DOCS = [
    ("1. Cover Letter", doc_cover),
    ("3. Promissory Note", doc_note),
    ("4. Personal Guarantee", doc_guarantee),
    ("5. Closing Instructions Letter", doc_closing),
    ("6. Insurance Request Letter", doc_insurance),
    ("7. Investor Information Sheet", doc_investor),
    ("8. Satisfaction and Release Request", doc_release),
]


def build_all(spec: dict, outdir: Path) -> list:
    outdir.mkdir(parents=True, exist_ok=True)
    expected = {f"{name}.docx" for name, _ in DOCS}
    written = []
    for name, fn in DOCS:
        path = outdir / f"{name}.docx"
        try:
            fn(spec, path)
            # The templates carry a NOTES FOR CLAUDE block that must never
            # reach a lender. Prove it rather than assume it.
            text = "\n".join(p.text for p in Document(path).paragraphs)
            if BANNED in text.upper():
                raise AssertionError("internal notes leaked into the output")
            written.append(path)
        except Exception as exc:                       # degrade, never fail
            logger.warning("%s failed: %s", name, exc)
    for stale in outdir.glob("*.docx"):
        if stale.name not in expected:
            stale.unlink()
            logger.info("Removed stale %s", stale.name)
    return written


def main(argv=None):
    ap = argparse.ArgumentParser(description="Lender document set")
    ap.add_argument("--spec", required=True)
    ap.add_argument("--outdir")
    args = ap.parse_args(argv)
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
    spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
    outdir = Path(args.outdir or
                  f"output/lender/{spec['deal_name'].replace(' ', '_')}")
    written = build_all(spec, outdir)
    for w in written:
        logger.info("Wrote %s", w.name)
    logger.info("%d of %d documents written to %s", len(written), len(DOCS), outdir)
    if len(written) != len(DOCS):
        logger.error("INCOMPLETE document set. %d failed, see warnings above.",
                     len(DOCS) - len(written))
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
